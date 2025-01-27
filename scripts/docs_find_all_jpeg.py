from ast import Index
import docx
import spacy
import openai
import os
import numpy as np
import boto3
from dotenv import load_dotenv
from pymilvus import (
    connections,
    FieldSchema,
    CollectionSchema,
    DataType,
    Collection,
    utility,
)
from docx import Document
from io import BytesIO, StringIO
from PIL import Image
import csv

# Загрузка переменных среды
load_dotenv("tokens.env")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY")
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY")
MINIO_BUCKET_NAME = os.getenv("MINIO_BUCKET_NAME")
MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT")
MINIO_REGION_NAME = os.getenv("MINIO_REGION_NAME")
MILVUS_HOST = os.getenv("MILVUS_HOST")
MILVUS_PORT = os.getenv("MILVUS_PORT")
MILVUS_COLLECTION = os.getenv("MILVUS_COLLECTION")

# Настройка важных переменных
name_of_collection_milvus = MILVUS_COLLECTION
name_of_bucket_minio = MINIO_BUCKET_NAME
path_of_doc_for_convert = r"C:\Project1\GITProjects\myproject2\pictures_test.docx"
openai.api_key = OPENAI_API_KEY

# Подключение к MinIO
s3_client = boto3.client(
    "s3",
    endpoint_url=MINIO_ENDPOINT,
    aws_access_key_id=MINIO_ACCESS_KEY,
    aws_secret_access_key=MINIO_SECRET_KEY,
    region_name=MINIO_REGION_NAME,
)

# Подключение к Milvus
connections.connect("default", host=MILVUS_HOST, port=MILVUS_PORT)

# Создание бакета MinIO, если он не существует
if s3_client.list_buckets().get("Buckets", None):
    existing_buckets = [
        bucket["Name"] for bucket in s3_client.list_buckets()["Buckets"]
    ]
    if name_of_bucket_minio not in existing_buckets:
        s3_client.create_bucket(Bucket=name_of_bucket_minio)

# Создание коллекции Milvus (если её нет)
collection_name = name_of_collection_milvus
if not utility.has_collection(collection_name):
    fields = [
        FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True),
        FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=1536),
        FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=65535),
        FieldSchema(name="reference", dtype=DataType.VARCHAR, max_length=65535),
    ]
    schema = CollectionSchema(fields, description="Коллекция для хранения данных")
    collection = Collection(name=collection_name, schema=schema)
else:
    collection = Collection(name=collection_name)

# Загрузка модели spaCy
nlp = spacy.load("ru_core_news_lg")


def create_embeddings(text):
    """Создает эмбеддинг текста с помощью OpenAI."""
    if not text.strip():
        return None
    response = openai.embeddings.create(input=[text], model="text-embedding-ada-002")
    return response.data[0].embedding


def split_text_logically(text):
    """Разделяет текст на логические блоки."""
    doc = nlp(text)
    logical_blocks = []
    current_block = []
    for sent in doc.sents:
        current_block.append(sent.text)
        if len(" ".join(current_block)) > 500:
            logical_blocks.append(" ".join(current_block))
            current_block = []
    if current_block:
        logical_blocks.append(" ".join(current_block))
    return logical_blocks


def save_table_to_minio(bucket_name, table_name, table_data):
    """Сохраняет таблицу в MinIO в формате CSV."""
    csv_buffer = StringIO()
    writer = csv.writer(csv_buffer)
    for row_data in table_data:
        writer.writerow(row_data)
    s3_client.put_object(Bucket=bucket_name, Key=table_name, Body=csv_buffer.getvalue())


def save_image_to_minio(bucket_name, image_name, image_data):
    """Сохраняет изображение в MinIO как JPEG файл и возвращает имя файла."""
    buffer = BytesIO()
    image_data = image_data.convert(
        "RGB"
    )  # Конвертируем изображение в RGB перед сохранением
    image_data.save(buffer, format="JPEG")
    s3_client.put_object(Bucket=bucket_name, Key=image_name, Body=buffer.getvalue())
    print(f"Изображение загружено в MinIO как {image_name}")
    return image_name  # Возвращаем имя файла вместо ссылки


def extract_content_from_word(word_path, bucket_name):
    """Извлекает текст, таблицы и изображения из Word файла, избегая дубликатов."""
    doc = Document(word_path)
    text_blocks_with_refs = []
    current_text_block = []
    current_table_data = []
    table_counter, image_counter = 1, 1
    last_was_table = False
    saved_images = set()  # Набор для отслеживания сохраненных изображений

    # Обработка текста и таблиц
    for idx, block in enumerate(doc.element.body):
        if block.tag.endswith("p"):  # Обработка параграфов
            paragraph = block.text.strip()
            if paragraph:
                if last_was_table and current_table_data:
                    table_name = f"table_{table_counter}.csv"
                    save_table_to_minio(bucket_name, table_name, current_table_data)
                    explanation = current_text_block[-1] if current_text_block else ""
                    text_blocks_with_refs.append(
                        {"text": explanation, "reference": table_name}
                    )
                    current_table_data = []
                    table_counter += 1
                current_text_block.append(paragraph)
                last_was_table = False

        elif block.tag.endswith("tbl"):  # Обработка таблиц
            table = next(t for t in doc.tables if t._tbl == block)
            current_table_data = [
                [cell.text.strip() for cell in row.cells] for row in table.rows
            ]
            # Обработка изображений внутри таблиц
            for row in table.rows:
                for cell in row.cells:
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        for run in paragraph.runs:
                            if run.element.xpath(".//a:blip"):
                                blip = run.element.xpath(".//a:blip")[0]
                                image_part = doc.part.related_parts[
                                    blip.get(
                                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                    )
                                ]
                                # Проверка на дубликат
                                if image_part in saved_images:
                                    continue
                                saved_images.add(
                                    image_part
                                )  # Добавляем изображение в набор

                                image_data = Image.open(BytesIO(image_part.blob))
                                image_name = f"image_{image_counter}.jpeg"
                                save_image_to_minio(bucket_name, image_name, image_data)

                                # Получаем следующий параграф для описания, если он существует
                                if paragraph_index + 1 < len(cell.paragraphs):
                                    text_after_image = cell.paragraphs[
                                        paragraph_index + 1
                                    ].text.strip()
                                else:
                                    text_after_image = "Описание отсутствует"

                                text_blocks_with_refs.append(
                                    {
                                        "text": text_after_image,
                                        "reference": image_name,
                                    }
                                )
                                print(
                                    f"Изображение {image_name} загружено с описанием: {text_after_image}"
                                )
                                image_counter += 1
            last_was_table = True

    # Обработка изображений вне таблиц
    paragraphs = iter(doc.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                # Найдено изображение
                blip = run.element.xpath(".//a:blip")[0]
                image_part = doc.part.related_parts[
                    blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                ]
                # Проверка на дубликат
                if image_part in saved_images:
                    continue
                saved_images.add(image_part)  # Добавляем изображение в набор

                image_data = Image.open(BytesIO(image_part.blob))
                image_name = f"image_{image_counter}.jpeg"
                save_image_to_minio(bucket_name, image_name, image_data)

                # Ищем текст непосредственно после изображения
                try:
                    next_paragraph = next(paragraphs)
                    text_after_image = (
                        next_paragraph.text.strip()
                        if next_paragraph.text.strip()
                        else "Описание отсутствует"
                    )
                except StopIteration:
                    text_after_image = "Описание отсутствует"

                text_blocks_with_refs.append(
                    {"text": text_after_image, "reference": image_name}
                )
                print(
                    f"Изображение {image_name} загружено с описанием: {text_after_image}"
                )
                image_counter += 1

    return text_blocks_with_refs, " ".join(current_text_block)


def process_content_from_word(word_path, bucket_name):
    """Обрабатывает текст, таблицы и изображения из Word файла и сохраняет в Milvus."""
    successful_embeddings_count = 0
    text_blocks_with_refs, full_text = extract_content_from_word(word_path, bucket_name)
    text_blocks = split_text_logically(full_text)

    for block in text_blocks:
        embedding = create_embeddings(block)
        if embedding is None:
            continue
        embedding_np = np.array(embedding, dtype=np.float32).tolist()
        data = [[embedding_np], [block], [""]]
        collection.insert(data)
        successful_embeddings_count += 1
        print(
            f"Эмбеддинг и текст успешно добавлены для блока {successful_embeddings_count}."
        )

    for ref_info in text_blocks_with_refs:
        text = ref_info["text"]
        reference = ref_info["reference"]
        embedding = create_embeddings(text)
        if embedding is None:
            continue
        embedding_np = np.array(embedding, dtype=np.float32).tolist()
        data = [[embedding_np], [text], [reference]]
        collection.insert(data)
        successful_embeddings_count += 1
        print(f"Эмбеддинг и пояснение успешно добавлены для объекта: {reference}")

    collection.flush()
    print("Все эмбеддинги и данные успешно добавлены в Milvus.")
    print(f"Количество успешно созданных эмбеддингов: {successful_embeddings_count}")


# Пример использования
word_path = path_of_doc_for_convert
process_content_from_word(word_path, name_of_bucket_minio)

# Создание и загрузка индекса в Milvus
index_params = {"index_type": "IVF_FLAT", "metric_type": "L2", "params": {"nlist": 128}}
collection.create_index(field_name="embedding", index_params=index_params)
collection.load()

print(f"Индекс успешно создан и коллекция '{collection_name}' загружена.")
