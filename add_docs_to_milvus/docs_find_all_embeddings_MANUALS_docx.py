import logging
import asyncio
from ast import Index
import re
import threading
import time
from venv import logger
import spacy
import openai
import os
import numpy as np
import boto3
from dotenv import load_dotenv
from pymilvus import (
    connections,
    FieldSchema,
    CollectionSchema,
    DataType,
    Collection,
    utility,
)
from docx import Document
from io import BytesIO
from PIL import Image
import tiktoken
from openpyxl import Workbook
import aioboto3
import aiofiles
from concurrent.futures import ThreadPoolExecutor
from concurrent.futures import ProcessPoolExecutor
from threading import Lock
import shutil
import sys
import requests
import http.client

# Загрузка переменных среды
load_dotenv("all_tockens.env")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # API токен OpenAI

MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY")  # Логин для подключенияMiniO
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY")  # Пароль для подключения MiniO
MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT")  # IP и порт MiniO
MINIO_REGION_NAME = os.getenv("MINIO_REGION_NAME")  # Регион MiniO
MINIO_BUCKET_NAME = os.getenv("MINIO_BUCKET_NAME")  # Название Бакета MiniO
MINIO_FOLDER_DOCS_NAME_SPRAVOCHNIK = os.getenv(
    "MINIO_FOLDER_DOCS_NAME_SPRAVOCHNIK"
)  # Название Папки хранения Таблиц/Изображений Справочника инженеров
MINIO_FOLDER_DOCS_NAME_MANUAL = os.getenv(
    "MINIO_FOLDER_DOCS_NAME_MANUAL"
)  # Название Папки хранения Таблиц/Изображений Мануала
MILVUS_DB_NAME_FIRST = os.getenv(
    "MILVUS_DB_NAME_FIRST"
)  # БД коллекций Милвуса(БД) с справочником

MILVUS_COLLECTION = os.getenv("MILVUS_COLLECTION")  # Коллекция Милвуса(БД)
MILVUS_HOST = os.getenv("MILVUS_HOST")  # IP Милвуса(БД)
MILVUS_PORT = os.getenv("MILVUS_PORT")  # Порт Милвуса(БД)

# Настройка логирования
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO
)

# =======================================================================================================

DOCX_DIRECTORY = (
    r"C:\Project1\GITProjects\scripts\Шлюхи"  # <================= Путь к файлам docx
)

end_name_docs = ".docx"  # <============ Конец имени исходного файла, названия коллекции

# =======================================================================================================

docx_files = [file for file in os.listdir(DOCX_DIRECTORY) if file.endswith(".docx")]
docx_count = len(docx_files)
print(f"Количество релевантных документов: {docx_count}")

# Настройка важных переменных
change_db_of_milvus = MILVUS_DB_NAME_FIRST  # <================================= Выбери бд, в которую будет записываться инфа (Справочник)
if not docx_files:
    raise ValueError("Нет файлов .docx в указанной директории.")

minio_folder_docs_name = MINIO_FOLDER_DOCS_NAME_MANUAL  # <================================= Выбери папку, в которую будет записываться инфа (Справочник)

name_of_bucket_minio = MINIO_BUCKET_NAME

milvus_collection_base = "Docs"  # <============== Название коллекции в Milvus


# name_documents = "Simrad Autopilot System AP70, AP80 Installation Manual"  # <============== Описание коллекции milvus

# path_of_doc_for_convert = r"C:\Project1\GITProjects\myproject2\add_docs_to_milvus\Simrad Autopilot System AP70, AP80 Installation Manual.docx"  # <============== Путь к файлу для добавления его в БД
# description_milvus_collection = name_documents + ".pdf"


openai.api_key = OPENAI_API_KEY

# Подключение к MinIO
s3_client = boto3.client(
    "s3",
    endpoint_url=MINIO_ENDPOINT,
    aws_access_key_id=MINIO_ACCESS_KEY,
    aws_secret_access_key=MINIO_SECRET_KEY,
    region_name=MINIO_REGION_NAME,
)
print(f'Логин "{MINIO_ACCESS_KEY}" для БД MiniO')  # Проверка LOG
print(f'Пароль "{MINIO_SECRET_KEY}" для БД MiniO')  # Проверка PSWD

# Подключение к Milvus
connections.connect(
    alias="default", host=MILVUS_HOST, port=MILVUS_PORT, db_name=change_db_of_milvus
)


count_collection_ends = 8


# Функция обрабатывает данные из Word, создает эмбеддинги и сохраняет все в Milvus
def process_content_from_word(
    word_path, bucket_name, description_milvus_collection, collection
):
    """Обрабатывает текст, таблицы и изображения из Word файла и сохраняет в Milvus."""
    successful_embeddings_count = 0
    text_blocks_with_refs, full_text = extract_content_from_word(
        word_path, bucket_name, description_milvus_collection
    )
    text_blocks = split_text_logically(full_text)

    # Данные для батч-вставки
    embeddings_batch = []
    texts_batch = []
    references_batch = []
    figure_ids_batch = []
    related_tables_batch = []
    manual_id_batch = []

    # **Добавление имени коллекции в Milvus как первый элемент**
    collection_name_block = description_milvus_collection  # Имя коллекции
    collection_name_embedding = create_embeddings(
        collection_name_block, description_milvus_collection
    )
    if collection_name_embedding:
        collection_name_embedding_np = np.array(
            collection_name_embedding, dtype=np.float32
        ).tolist()
        data = [
            [collection_name_embedding_np],
            [collection_name_block],
            [""],
            [""],
            [""],
            [description_milvus_collection],
        ]
        collection.insert(data)
        successful_embeddings_count += 1

    # Обрабатываем текстовые блоки
    for block in text_blocks:
        if block and block.strip():  # Проверка, чтобы блок текста не был пустым
            embedding = create_embeddings(block, description_milvus_collection)
            if embedding is None:
                continue
            embedding_np = np.array(embedding, dtype=np.float32).tolist()
            data = [
                [embedding_np],
                [block],
                [""],
                [""],
                [""],
                [description_milvus_collection],
            ]
            collection.insert(data)
            successful_embeddings_count += 1

    # Обрабатываем метаинформацию (ссылки, рисунки, таблицы)
    for ref_info in text_blocks_with_refs:
        text = ref_info["text"]
        reference = ref_info["reference"]
        figure_id = ref_info["figure_id"]
        related_table = ref_info["related_table"]
        manual_id = ref_info["manual_id"]

        if text and text.strip():
            embedding = create_embeddings(text, description_milvus_collection)
            if embedding is None:
                continue
            embedding_np = np.array(embedding, dtype=np.float32).tolist()
            data = [
                [embedding_np],
                [text],
                [reference],
                [figure_id],
                [related_table],
                [manual_id],
            ]
            collection.insert(data)
            successful_embeddings_count += 1

    collection.flush()
    # print(f"Количество успешно созданных эмбеддингов: {successful_embeddings_count}")


def insert_batch_to_milvus(
    embeddings, texts, references, figure_ids, related_tables, collection
):
    """Вставляет батч данных в Milvus."""
    try:
        data = [embeddings, texts, references, figure_ids, related_tables]
        collection.insert(data)
        print(
            f"Батч из {len(embeddings)} записей успешно добавлен в  коллекцию: {collection.name}"
        )
    except Exception as e:
        print(f"Ошибка при вставке данных в Milvus: {e}")


# Функция создает эмбеддинги ко всему тексту (описание рисунков, текста таблиц, любого текста)
def create_embeddings(
    text, description_milvus_collection, max_retries=5, retry_delay=5
):
    """
    Создает эмбеддинг текста с помощью OpenAI с повторными попытками в случае ошибки.

    Args:
        text (str): Текст для создания эмбеддинга.
        max_retries (int): Максимальное количество попыток перед завершением.
        retry_delay (int): Задержка между попытками (в секундах).

    Returns:
        list: Эмбеддинг текста или завершает скрипт при неудачных попытках.
    """
    if not text.strip():
        return None

    attempt = 0  # Счетчик попыток

    while attempt < max_retries:
        try:
            response = openai.embeddings.create(
                input=[text], model="text-embedding-ada-002"
            )
            return response.data[0].embedding  # Возврат успешного эмбеддинга

        except Exception as e:
            attempt += 1
            print(
                f"Попытка {attempt}/{max_retries} коллекции {description_milvus_collection}: Ошибка при создании эмбеддинга: {e}"
            )

            # Если ошибка связана с API ограничениями (например, 429 или 500)
            if "rate limit" in str(e).lower() or "server error" in str(e).lower():
                print(f"Пауза {retry_delay} секунд перед повторной попыткой...")
                time.sleep(retry_delay)
                continue

            # Если ошибка связана с неподдерживаемым регионом или другой критической причиной
            if "unsupported_country_region_territory" in str(e):
                print("Критическая ошибка: Неподдерживаемый регион.")
                break  # Завершить попытки

            # Задержка перед следующей попыткой для других ошибок
            time.sleep(retry_delay)

    # Если все попытки не удались, завершить выполнение скрипта
    print(
        f"Не удалось создать эмбеддинг после {max_retries} попыток. Завершение скрипта."
    )
    send_telegram_error_message()


def split_text_logically(text):
    """
    Разделяет текст на логические блоки по 500 символов, соблюдая границы абзацев.
    """
    paragraphs = text.split("\n")  # Разделяем текст на абзацы
    logical_blocks = []  # Список для хранения логических блоков
    current_block = ""  # Текущий блок текста

    for paragraph in paragraphs:
        paragraph = paragraph.strip()  # Убираем лишние пробелы
        if not paragraph:
            continue  # Пропускаем пустые абзацы

        # Если добавление текущего абзаца не превышает 500 символов, добавляем его
        if len(current_block) + len(paragraph) + 1 <= 500:  # +1 для пробела/разделителя
            current_block += paragraph + " "
        else:
            # Если текущий блок превышает 500 символов, сохраняем его и начинаем новый
            logical_blocks.append(current_block.strip())
            current_block = paragraph + " "

    # Добавляем последний блок, если он не пустой
    if current_block.strip():
        logical_blocks.append(current_block.strip())

    return logical_blocks


def split_table_text_logically(table_data):
    """
    Разделяет таблицу на логические блоки, обрабатывая каждую строку индивидуально.

    Args:
        table_data (list of list of str): Данные таблицы в виде списка строк, где каждая строка - это список ячеек.

    Returns:
        list of str: Список текстовых строк таблицы.
    """
    logical_blocks = []

    for row in table_data:
        # Объединяем ячейки строки через табуляцию
        row_text = "\t".join(row)
        logical_blocks.append(row_text)  # Добавляем строку как отдельный блок

    return logical_blocks


# Функция сохраняет исходный файл в MiniO
def save_table_to_minio(bucket_name, description_milvus_collection):
    """
    Сохраняет файл, соответствующий значению переменной description_milvus_collection, в MinIO.

    Args:
        bucket_name (str): Название бакета в MinIO.
        description_milvus_collection (str): Имя файла для поиска и сохранения.
    """
    try:
        # Генерируем полный путь к файлу
        file_path = os.path.join(DOCX_DIRECTORY, description_milvus_collection)

        # Проверяем, существует ли файл
        if not os.path.exists(file_path):
            raise FileNotFoundError(
                f"Файл {description_milvus_collection} не найден в {DOCX_DIRECTORY}."
            )

        # Открываем файл и читаем его содержимое
        with open(file_path, "rb") as file:
            file_data = file.read()

        # Генерируем ключ для сохранения в MinIO
        minio_key = f"{minio_folder_docs_name}/{description_milvus_collection}"

        # Сохраняем файл в MinIO
        s3_client.put_object(
            Bucket=bucket_name,
            Key=minio_key,
            Body=file_data,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ContentDisposition="inline",  # Указывает браузеру открывать файл, а не скачивать
        )

        logger.info(
            f"Файл {description_milvus_collection} успешно сохранён в MinIO под ключом {minio_key}."
        )

    except Exception as e:
        logger.info(
            f"Ошибка при сохранении файла {description_milvus_collection} в MinIO: {e}"
        )


# Функция обрабатывает Word документ, извлекая таблицы, текст, изображения из документа и сохраняя в MiniO
def extract_content_from_word(word_path, bucket_name, description_milvus_collection):
    """Извлекает текст, таблицы и изображения из Word файла, избегая дубликатов."""
    doc = Document(word_path)
    text_blocks_with_refs = []
    current_text_block = []
    current_table_data = []
    table_counter, image_counter = 1, 1
    last_was_table = False
    saved_images = set()  # Набор для отслеживания сохраненных изображений

    # Обработка текста и таблиц
    for idx, block in enumerate(doc.element.body):
        if block.tag.endswith("p"):  # Обработка параграфов
            paragraph = block.text.strip()
            if paragraph:
                if last_was_table and current_table_data:
                    # Сохраняем текущую собранную таблицу в MinIO как одну таблицу
                    table_name = f"table_{table_counter}"
                    table_name_xlsx = f"{table_name}.xlsx"

                    # Сохраняем описание таблицы
                    explanation = current_text_block[-1] if current_text_block else ""

                    text_blocks_with_refs.append(
                        {
                            "text": explanation,
                            "reference": "",
                            "figure_id": "",
                            "related_table": "",
                            "manual_id": description_milvus_collection,
                        }
                    )
                    # Обрабатываем текст из таблицы и сохраняем в Milvus
                    table_text_blocks = split_table_text_logically(current_table_data)
                    for block in table_text_blocks:
                        text_blocks_with_refs.append(
                            {
                                "text": block,
                                "reference": "",
                                "figure_id": "",
                                "related_table": table_name_xlsx,
                                "manual_id": description_milvus_collection,
                            }
                        )
                    current_table_data = []  # Сброс текущих данных таблицы
                    table_counter += 1
                current_text_block.append(paragraph)
                last_was_table = False

        elif block.tag.endswith("tbl"):  # Обработка таблиц
            table = next(t for t in doc.tables if t._tbl == block)
            table_data = [
                [cell.text.strip() for cell in row.cells] for row in table.rows
            ]

            # Объединяем таблицы, если они идут подряд
            if last_was_table:
                current_table_data.extend(table_data)
            else:
                current_table_data = table_data

            # Обработка изображений внутри таблиц
            for row in table.rows:
                for cell in row.cells:
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        for run in paragraph.runs:
                            if run.element.xpath(".//a:blip"):
                                blip = run.element.xpath(".//a:blip")[0]
                                image_part = doc.part.related_parts[
                                    blip.get(
                                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                    )
                                ]
                                # Проверка на дубликат
                                if image_part in saved_images:
                                    continue
                                saved_images.add(image_part)

                                image_data = Image.open(BytesIO(image_part.blob))
                                image_name = f"image_{image_counter}.jpeg"

                                # Получаем следующий параграф для описания, если он существует
                                if paragraph_index + 1 < len(cell.paragraphs):
                                    text_after_image = cell.paragraphs[
                                        paragraph_index + 1
                                    ].text.strip()
                                else:
                                    text_after_image = "Описание отсутствует"

                                # Добавляем запись с `text`, `reference` и `figure_id`
                                text_blocks_with_refs.append(
                                    {
                                        "text": text_after_image,
                                        "reference": "",
                                        "figure_id": "",
                                        "related_table": f"table_{table_counter}.xlsx",  # Смещение на 1 для таблицы
                                        "manual_id": description_milvus_collection,
                                    }
                                )
                                print(
                                    f"Изображение {image_name} загружено с описанием: {text_after_image}, related_table: {table_name_xlsx}"
                                )
                                image_counter += 1
            last_was_table = True

    # Обработка изображений вне таблиц
    paragraphs = iter(doc.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                # Найдено изображение
                blip = run.element.xpath(".//a:blip")[0]
                image_part = doc.part.related_parts[
                    blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                ]
                # Проверка на дубликат
                if image_part in saved_images:
                    continue
                saved_images.add(image_part)

                image_data = Image.open(BytesIO(image_part.blob))
                image_name = f"image_{image_counter}.jpeg"

                # Ищем текст непосредственно после изображения
                try:
                    next_paragraph = next(paragraphs)
                    text_after_image = (
                        next_paragraph.text.strip()
                        if next_paragraph.text.strip()
                        else "Описание отсутствует"
                    )
                except StopIteration:
                    text_after_image = "Описание отсутствует"

                text_blocks_with_refs.append(
                    {
                        "text": text_after_image,
                        "reference": "",
                        "figure_id": "",
                        "related_table": "",  # Поле остаётся пустым
                        "manual_id": description_milvus_collection,
                    }
                )
                print(
                    f"Изображение {image_name} загружено с описанием: {text_after_image}"
                )
                image_counter += 1

    return text_blocks_with_refs, " ".join(current_text_block)


collection_lock = threading.Lock()


def send_telegram_error_message():
    """
    Отправляет сообщение об ошибке в Telegram перед завершением программы.

    Args:
        bot_token (str): Токен Telegram-бота.
        chat_id (str): Идентификатор чата (или пользователя), куда отправить сообщение.
        error_message (str): Сообщение об ошибке для отправки.

    Raises:
        Exception: Если возникла ошибка при отправке сообщения.
    """
    try:
        conn = http.client.HTTPSConnection("api.telegram.org")

        payload = (
            '\n{\n  "chat_id": "5746497552",\n  "text": "Ссаный впн упал"\n}'.encode(
                "utf-8"
            )
        )

        headers = {"Content-Type": "application/json"}

        conn.request(
            "POST",
            "/bot7219050865:AAFuYYrlMdyNTOd2Ffy83sFY-byESBF7hwQ/sendMessage",
            payload,
            headers,
        )

        res = conn.getresponse()
        data = res.read()

        print(data.decode("utf-8"))

    except Exception as e:
        print(f"Не удалось отправить сообщение в Telegram: {e}")


def send_telegram_complite_message():
    """
    Отправляет сообщение об ошибке в Telegram перед завершением программы.

    Args:
        bot_token (str): Токен Telegram-бота.
        chat_id (str): Идентификатор чата (или пользователя), куда отправить сообщение.
        error_message (str): Сообщение об ошибке для отправки.

    Raises:
        Exception: Если возникла ошибка при отправке сообщения.
    """
    try:
        conn = http.client.HTTPSConnection("api.telegram.org")

        payload = '\n{\n  "chat_id": "5746497552",\n  "text": "Все документы загружены"\n}'.encode(
            "utf-8"
        )

        headers = {"Content-Type": "application/json"}

        conn.request(
            "POST",
            "/bot7219050865:AAFuYYrlMdyNTOd2Ffy83sFY-byESBF7hwQ/sendMessage",
            payload,
            headers,
        )

        res = conn.getresponse()
        data = res.read()

        print(data.decode("utf-8"))

    except Exception as e:
        print(f"Не удалось отправить сообщение в Telegram: {e}")


def get_unique_collection_name(base_name, start_index):
    """
    Получает уникальное имя коллекции, проверяя существование коллекций в Milvus.

    Args:
        base_name (str): Базовое имя коллекции (например, "Docs").
        start_index (int): Начальный индекс для генерации имени коллекции.

    Returns:
        str: Уникальное имя коллекции.
    """
    global count_collection_ends  # Указываем, что переменная глобальная

    with collection_lock:  # Только один поток заходит сюда одновременно
        while True:

            collection_name = f"{base_name}{start_index}"

            # print(f"Проверяем коллекцию: {collection_name}")
            try:
                if not utility.has_collection(collection_name):
                    # Если коллекции нет, создаём её
                    if start_index % 20 == 0:
                        count_collection_ends = start_index
                    return collection_name

            except Exception as e:
                print(f"Ошибка при проверке коллекции {collection_name}: {e}")
            start_index += 1

            # print(f"Увеличиваем индекс: {start_index}")


# метод для перемещения отработанных мануалов(оригиналов)
def move_file(file_name, destination_path):
    """
    Перемещает файл из текущего местоположения в указанный путь.

    Args:
        file_name (str): Название файла для перемещения.
        destination_path (str): Путь, куда переместить файл.

    Raises:
        FileNotFoundError: Если файл не найден.
        Exception: Если возникает ошибка при перемещении.
    """
    try:
        # Получаем полный путь к файлу
        current_directory = destination_path  # Текущая директория
        source_path = os.path.join(current_directory, file_name)

        # Проверяем, существует ли файл
        if not os.path.exists(source_path):
            raise FileNotFoundError(
                f"Файл {file_name} не найден в {current_directory}."
            )

        # Проверяем, существует ли целевая папка, и создаем, если нет
        if not os.path.exists(f"{destination_path}\\ready"):
            os.makedirs(f"{destination_path}\\ready")

        # Полный путь к новому местоположению файла
        target_path = os.path.join(f"{destination_path}\\ready", file_name)

        # Перемещаем файл
        shutil.move(source_path, target_path)
        print(f"Файл {file_name} успешно перемещен в {destination_path}\\ready")

    except Exception as e:
        print(f"Ошибка при перемещении файла {file_name}: {e}")


def process_docx_file(docx_file, s3_client, path_to_save_manuals):
    """Асинхронная функция для обработки одного файла."""
    print(f"Метод process_docx_file запустился для {docx_file}")

    # Работа с Milvus

    name_documents = os.path.splitext(docx_file)[0]
    path_of_doc_for_convert = os.path.join(DOCX_DIRECTORY, docx_file)
    description_milvus_collection = name_documents + end_name_docs
    # print(f"description_milvus_collection {description_milvus_collection}")

    # Уникальное имя коллекции
    """milvus_collection = get_unique_collection_name(
        milvus_collection_base, count_collection_ends
    )"""
    milvus_collection = "Manuals"

    # Загрузка файла в MinIO
    minio_key = f"{minio_folder_docs_name}/{description_milvus_collection}"
    # print(f"minio_key {minio_key}")
    # bucket_name = MINIO_BUCKET_NAME
    """try:
        with open(f"path_to_file/{docx_file}", "rb") as file_data:
            s3_client.put_object(Bucket=bucket_name, Key=minio_key, Body=file_data)
    except Exception as e:
        print(f"Ошибка загрузки в MinIO: {e}")"""

    if not utility.has_collection(milvus_collection):
        fields = [
            FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True),
            FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=1536),
            FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=65535),
            FieldSchema(name="reference", dtype=DataType.VARCHAR, max_length=65535),
            FieldSchema(name="figure_id", dtype=DataType.VARCHAR, max_length=100),
            FieldSchema(name="related_table", dtype=DataType.VARCHAR, max_length=65535),
            FieldSchema(name="manual_id", dtype=DataType.VARCHAR, max_length=256),
        ]
        schema = CollectionSchema(fields, description="Коллекция со всеми мануалами")
        collection = Collection(name=milvus_collection, schema=schema)
    else:
        collection = Collection(name=milvus_collection)

    process_content_from_word(
        path_of_doc_for_convert,
        name_of_bucket_minio,
        description_milvus_collection,
        collection,
    )

    print(
        "---------------------------------------------------------------------------------------------------"
    )
    print("Начало загрузки коллекции в Milvus")
    # Создание и загрузка индекса в Milvus
    index_params = {
        "index_type": "IVF_FLAT",
        "metric_type": "L2",
        "params": {"nlist": 128},
    }
    collection.create_index(field_name="embedding", index_params=index_params)
    collection.load()
    print(
        f"Конец загрузки коллекции в Milvus {description_milvus_collection}, начало загрузки в MiniO"
    )
    save_table_to_minio(name_of_bucket_minio, description_milvus_collection)
    move_file(description_milvus_collection, path_to_save_manuals)

    print(
        "---------------------------------------------------------------------------------------------------"
    )
    print(
        f"Индекс успешно создан и коллекция '{milvus_collection}''{description_milvus_collection}' загружена в БД '{change_db_of_milvus}'"
    )
    print(
        "---------------------------------------------------------------------------------------------------"
    )


def main():
    # Создаем подключение один раз
    try:
        milvus_collection = connections.connect(
            alias="default",
            host=MILVUS_HOST,
            port=MILVUS_PORT,
            db_name=change_db_of_milvus,
        )
        print("Подключение к Milvus успешно установлено!")
    except Exception as e:
        print(f"Ошибка подключения к Milvus: {e}")
    s3_client = boto3.client(
        "s3",
        endpoint_url=MINIO_ENDPOINT,
        aws_access_key_id=MINIO_ACCESS_KEY,
        aws_secret_access_key=MINIO_SECRET_KEY,
        region_name=MINIO_REGION_NAME,
    )

    # Передаем подключение в потоки
    with ThreadPoolExecutor(max_workers=1) as executor:
        executor.map(
            lambda docx_file: process_docx_file(docx_file, s3_client, DOCX_DIRECTORY),
            docx_files,
        )
    send_telegram_complite_message()


if __name__ == "__main__":
    main()


print(f"Все коллекции загружены.")
