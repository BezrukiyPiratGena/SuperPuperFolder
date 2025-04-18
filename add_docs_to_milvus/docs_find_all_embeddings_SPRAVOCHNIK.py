from ast import Index
from ctypes import alignment
from openpyxl.styles import Border, Side, Alignment
import re
import time
import uuid
from venv import logger
import spacy
import openai
import os
import numpy as np
import boto3
from dotenv import load_dotenv
from pymilvus import (
    connections,
    FieldSchema,
    CollectionSchema,
    DataType,
    Collection,
    utility,
)
from docx import Document
from io import BytesIO
from PIL import Image
import tiktoken
from openpyxl import Workbook
from openpyxl.styles import Font

# –ó–∞–≥—Ä—É–∑–∫–∞ –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö —Å—Ä–µ–¥—ã
load_dotenv("all_tockens.env")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # API —Ç–æ–∫–µ–Ω OpenAI

MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY")  # –õ–æ–≥–∏–Ω –¥–ª—è –ø–æ–¥–∫–ª—é—á–µ–Ω–∏—èMiniO
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY")  # –ü–∞—Ä–æ–ª—å –¥–ª—è –ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è MiniO
MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT")  # IP –∏ –ø–æ—Ä—Ç MiniO
MINIO_REGION_NAME = os.getenv("MINIO_REGION_NAME")  # –†–µ–≥–∏–æ–Ω MiniO
MINIO_BUCKET_NAME = os.getenv("MINIO_BUCKET_NAME")  # –ù–∞–∑–≤–∞–Ω–∏–µ –ë–∞–∫–µ—Ç–∞ MiniO
MINIO_FOLDER_DOCS_NAME_SPRAVOCHNIK = os.getenv(
    "MINIO_FOLDER_DOCS_NAME_SPRAVOCHNIK"
)  # –ù–∞–∑–≤–∞–Ω–∏–µ –ü–∞–ø–∫–∏ —Ö—Ä–∞–Ω–µ–Ω–∏—è –¢–∞–±–ª–∏—Ü/–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –°–ø—Ä–∞–≤–æ—á–Ω–∏–∫–∞ –∏–Ω–∂–µ–Ω–µ—Ä–æ–≤
MILVUS_DB_NAME_FIRST = os.getenv(
    "MILVUS_DB_NAME_FIRST"
)  # –ë–î –∫–æ–ª–ª–µ–∫—Ü–∏–π –ú–∏–ª–≤—É—Å–∞(–ë–î) —Å —Å–ø—Ä–∞–≤–æ—á–Ω–∏–∫–æ–º

MILVUS_COLLECTION = os.getenv("MILVUS_COLLECTION")  # –ö–æ–ª–ª–µ–∫—Ü–∏—è –ú–∏–ª–≤—É—Å–∞(–ë–î)
MILVUS_HOST = os.getenv("MILVUS_HOST")  # IP –ú–∏–ª–≤—É—Å–∞(–ë–î)
MILVUS_PORT = os.getenv("MILVUS_PORT")  # –ü–æ—Ä—Ç –ú–∏–ª–≤—É—Å–∞(–ë–î)
MILVUS_USER = os.getenv("MILVUS_USER")  # –õ–æ–≥–∏–Ω –ú–∏–ª–≤—É—Å–∞(–ë–î)
MILVUS_PASSWORD = os.getenv("MILVUS_PASSWORD")  # –ü–∞—Ä–æ–ª—å –ú–∏–ª–≤—É—Å–∞(–ë–î)

# –ù–∞—Å—Ç—Ä–æ–π–∫–∞ –≤–∞–∂–Ω—ã—Ö –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö
change_db_of_milvus = MILVUS_DB_NAME_FIRST  # <================================= –í—ã–±–µ—Ä–∏ –±–¥, –≤ –∫–æ—Ç–æ—Ä—É—é –±—É–¥–µ—Ç –∑–∞–ø–∏—Å—ã–≤–∞—Ç—å—Å—è –∏–Ω—Ñ–∞ (–°–ø—Ä–∞–≤–æ—á–Ω–∏–∫)

name_of_collection_milvus = MILVUS_COLLECTION

minio_folder_docs_name = MINIO_FOLDER_DOCS_NAME_SPRAVOCHNIK  # <================================= –í—ã–±–µ—Ä–∏ –ø–∞–ø–∫—É, –≤ –∫–æ—Ç–æ—Ä—É—é –±—É–¥–µ—Ç –∑–∞–ø–∏—Å—ã–≤–∞—Ç—å—Å—è –∏–Ω—Ñ–∞ (–°–ø—Ä–∞–≤–æ—á–Ω–∏–∫)

name_of_bucket_minio = MINIO_BUCKET_NAME
name_of_origin_doc = "test_docs.docx"  # <====================================================================== –ù–∞–∑–≤–∞–Ω–∏–µ —Ñ–∞–π–ª–∞ –¥–ª—è –¥–æ–±–∞–≤–ª–µ–Ω–∏—è –µ–≥–æ –≤ –ë–î
path_of_doc_for_convert = rf"C:\Project1\GITProjects\myproject2\add_docs_to_milvus\{name_of_origin_doc}"  # <============== –ü—É—Ç—å –∫ —Ñ–∞–π–ª—É –¥–ª—è –¥–æ–±–∞–≤–ª–µ–Ω–∏—è –µ–≥–æ –≤ –ë–î
description_milvus_collection = (
    "–°–ø—Ä–∞–≤–æ—á–Ω–∏–∫ –°–ò–†"  # <============== –û–ø–∏—Å–∞–Ω–∏–µ –∫–æ–ª–ª–µ–∫—Ü–∏–∏ milvus
)
openai.api_key = OPENAI_API_KEY

# –ü–æ–¥–∫–ª—é—á–µ–Ω–∏–µ –∫ MinIO
s3_client = boto3.client(
    "s3",
    endpoint_url=MINIO_ENDPOINT,
    aws_access_key_id=MINIO_ACCESS_KEY,
    aws_secret_access_key=MINIO_SECRET_KEY,
    region_name=MINIO_REGION_NAME,
)
print(f'–õ–æ–≥–∏–Ω "{MINIO_ACCESS_KEY}" –¥–ª—è –ë–î MiniO')  # –ü—Ä–æ–≤–µ—Ä–∫–∞ LOG
print(f'–ü–∞—Ä–æ–ª—å "{MINIO_SECRET_KEY}" –¥–ª—è –ë–î MiniO')  # –ü—Ä–æ–≤–µ—Ä–∫–∞ PSWD

# –ü–æ–¥–∫–ª—é—á–µ–Ω–∏–µ –∫ Milvus
connections.connect(
    alias="default",
    host=MILVUS_HOST,
    port=MILVUS_PORT,
    db_name=change_db_of_milvus,
    user=MILVUS_USER,
    password=MILVUS_PASSWORD,
)

# –°–æ–∑–¥–∞–Ω–∏–µ –±–∞–∫–µ—Ç–∞ MinIO, –µ—Å–ª–∏ –æ–Ω –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç
if s3_client.list_buckets().get("Buckets", None):
    existing_buckets = [
        bucket["Name"] for bucket in s3_client.list_buckets()["Buckets"]
    ]
    if name_of_bucket_minio not in existing_buckets:
        s3_client.create_bucket(Bucket=name_of_bucket_minio)

# –°–æ–∑–¥–∞–Ω–∏–µ –∫–æ–ª–ª–µ–∫—Ü–∏–∏ Milvus (–µ—Å–ª–∏ –µ—ë –Ω–µ—Ç)
collection_name = name_of_collection_milvus
if not utility.has_collection(collection_name):
    fields = [
        FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True),
        FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=1536),
        FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=65535),
        FieldSchema(name="reference", dtype=DataType.VARCHAR, max_length=65535),
        FieldSchema(name="figure_id", dtype=DataType.VARCHAR, max_length=100),
        FieldSchema(name="related_table", dtype=DataType.VARCHAR, max_length=65535),
        FieldSchema(name="origin_name_docs", dtype=DataType.VARCHAR, max_length=65535),
    ]
    schema = CollectionSchema(fields, description=description_milvus_collection)
    collection = Collection(name=collection_name, schema=schema)
else:
    collection = Collection(name=collection_name)

# –ó–∞–≥—Ä—É–∑–∫–∞ –º–æ–¥–µ–ª–∏ spaCy
nlp = spacy.load("ru_core_news_lg")

count_image_to_save = 1
count_table_to_save = 1
count_embedding_save = 1


# –§—É–Ω–∫—Ü–∏—è —Å–æ–∑–¥–∞–µ—Ç —ç–º–±–µ–¥–¥–∏–Ω–≥–∏ –∫–æ –≤—Å–µ–º—É —Ç–µ–∫—Å—Ç—É (–æ–ø–∏—Å–∞–Ω–∏–µ —Ä–∏—Å—É–Ω–∫–æ–≤, —Ç–µ–∫—Å—Ç–∞ —Ç–∞–±–ª–∏—Ü, –ª—é–±–æ–≥–æ —Ç–µ–∫—Å—Ç–∞)
def create_embeddings(text, pause_duration=2):
    """–°–æ–∑–¥–∞—ë—Ç —ç–º–±–µ–¥–¥–∏–Ω–≥ —Ç–µ–∫—Å—Ç–∞ —Å –ø–æ–º–æ—â—å—é OpenAI, –ø–æ–≤—Ç–æ—Ä—è—è –∑–∞–ø—Ä–æ—Å –¥–æ —É—Å–ø–µ—à–Ω–æ–≥–æ –æ—Ç–≤–µ—Ç–∞."""
    if not text.strip():
        return None

    while True:
        count_try = 1
        try:
            num_tokens = count_tokens(text)
            print(f"–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —Ç–æ–∫–µ–Ω–æ–≤ –≤ —Ç–µ–∫—Å—Ç–µ: {num_tokens}")
            response = openai.embeddings.create(
                input=[text], model="text-embedding-ada-002"
            )
            time.sleep(pause_duration)  # –ø–∞—É–∑–∞ –º–µ–∂–¥—É –∑–∞–ø—Ä–æ—Å–∞–º–∏
            return response.data[0].embedding

        except Exception as e:
            # –õ–æ–≥–∏—Ä—É–µ–º –æ—à–∏–±–∫—É –∏ –ø–æ–≤—Ç–æ—Ä—è–µ–º –∑–∞–ø—Ä–æ—Å
            print(
                f"‚ùó –û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ–∑–¥–∞–Ω–∏–∏ —ç–º–±–µ–¥–¥–∏–Ω–≥–∞: {e}. ({count_try})–ü–æ–≤—Ç–æ—Ä–Ω–∞—è –ø–æ–ø—ã—Ç–∫–∞ —á–µ—Ä–µ–∑ {pause_duration} —Å–µ–∫..."
            )
            time.sleep(pause_duration)
            count_try += 1
            # –∏ —Ü–∏–∫–ª –ø—Ä–æ–¥–æ–ª–∂–∏—Ç—Å—è, –ø–æ–∫–∞ –Ω–µ –≤–µ—Ä–Ω—ë—Ç—Å—è embedding


# –ü–æ–¥—Å—á–µ—Ç —Ç–æ–∫–µ–Ω–æ–≤ –∫–∞–∫–æ–≥–æ-—Ç–æ –æ—Ç—Ä—ã–≤–∫–∞ —Ç–µ–∫—Å—Ç–∞
def count_tokens(text, model="text-embedding-ada-002"):
    """
    –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ—Ç –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Ç–æ–∫–µ–Ω–æ–≤ –≤ —Ç–µ–∫—Å—Ç–µ –¥–ª—è —É–∫–∞–∑–∞–Ω–Ω–æ–π –º–æ–¥–µ–ª–∏ OpenAI.

    Args:
        text (str): –¢–µ–∫—Å—Ç, –¥–ª—è –∫–æ—Ç–æ—Ä–æ–≥–æ –Ω—É–∂–Ω–æ –ø–æ—Å—á–∏—Ç–∞—Ç—å —Ç–æ–∫–µ–Ω—ã.
        model (str): –ù–∞–∑–≤–∞–Ω–∏–µ –º–æ–¥–µ–ª–∏ OpenAI (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é text-embedding-ada-002).

    Returns:
        int: –ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —Ç–æ–∫–µ–Ω–æ–≤ –≤ —Ç–µ–∫—Å—Ç–µ.
    """
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    return len(tokens)


# –ü—Ä–∏–º–µ—Ä –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è
text = "–≠—Ç–æ –ø—Ä–∏–º–µ—Ä —Ç–µ–∫—Å—Ç–∞ –¥–ª—è –ø–æ–¥—Å—á–µ—Ç–∞ —Ç–æ–∫–µ–Ω–æ–≤."
num_tokens = count_tokens(text)
print(f"–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —Ç–æ–∫–µ–Ω–æ–≤: {num_tokens}")


# –§—É–Ω–∫—Ü–∏—è —Å–æ–∑–¥–∞–µ—Ç –ª–æ–≥ –±–ª–æ–∫–∏ –∏–∑ —Ç–µ–∫—Å—Ç–∞ –≤–Ω–µ —Ç–∞–±–ª–∏—Ü
def split_text_logically(text):
    """–†–∞–∑–¥–µ–ª—è–µ—Ç —Ç–µ–∫—Å—Ç –Ω–∞ –ª–æ–≥–∏—á–µ—Å–∫–∏–µ –±–ª–æ–∫–∏."""
    doc = nlp(text)
    logical_blocks = []
    current_block = []
    for sent in doc.sents:
        current_block.append(sent.text)
        if len(" ".join(current_block)) > 500:
            logical_blocks.append(" ".join(current_block))
            current_block = []
    if current_block:
        logical_blocks.append(" ".join(current_block))
    return logical_blocks


def split_table_text_logically(table_data):
    """
    –†–∞–∑–¥–µ–ª—è–µ—Ç —Ç–∞–±–ª–∏—Ü—É –Ω–∞ –ª–æ–≥–∏—á–µ—Å–∫–∏–µ –±–ª–æ–∫–∏, –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞—è –∫–∞–∂–¥—É—é —Å—Ç—Ä–æ–∫—É –∏–Ω–¥–∏–≤–∏–¥—É–∞–ª—å–Ω–æ.

    Args:
        table_data (list of list of str): –î–∞–Ω–Ω—ã–µ —Ç–∞–±–ª–∏—Ü—ã –≤ –≤–∏–¥–µ —Å–ø–∏—Å–∫–∞ —Å—Ç—Ä–æ–∫, –≥–¥–µ –∫–∞–∂–¥–∞—è —Å—Ç—Ä–æ–∫–∞ - —ç—Ç–æ —Å–ø–∏—Å–æ–∫ —è—á–µ–µ–∫.

    Returns:
        list of str: –°–ø–∏—Å–æ–∫ —Ç–µ–∫—Å—Ç–æ–≤—ã—Ö —Å—Ç—Ä–æ–∫ —Ç–∞–±–ª–∏—Ü—ã.
    """
    logical_blocks = []
    count = True

    for row in table_data:
        # –û–±—ä–µ–¥–∏–Ω—è–µ–º —è—á–µ–π–∫–∏ —Å—Ç—Ä–æ–∫–∏ —á–µ—Ä–µ–∑ —Ç–∞–±—É–ª—è—Ü–∏—é
        row_text = "\t".join(row)
        if (
            "–ö–æ–Ω–µ—Ü —Ç–∞–±–ª–∏—Ü—ã, –Ω–∞—á–∞–ª–æ –ø–æ—è—Å–Ω–µ–Ω–∏–π" in row
        ):  # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –≤—Å—Ç—Ä–µ—á–∞–µ—Ç—Å—è –ª–∏ 'Cir73SPb+' –≤ —Å—Ç—Ä–æ–∫–µ
            print(
                "üî¥ –û–±–Ω–∞—Ä—É–∂–µ–Ω '–ö–æ–Ω–µ—Ü —Ç–∞–±–ª–∏—Ü—ã, –Ω–∞—á–∞–ª–æ –ø–æ—è—Å–Ω–µ–Ω–∏–π'. –ü—Ä–µ—Ä—ã–≤–∞–µ–º –æ–±—Ä–∞–±–æ—Ç–∫—É —Ç–∞–±–ª–∏—Ü—ã."
            )
            count = False
        if count == False:
            continue
        else:
            logical_blocks.append(row_text)  # –î–æ–±–∞–≤–ª—è–µ–º —Å—Ç—Ä–æ–∫—É –∫–∞–∫ –æ—Ç–¥–µ–ª—å–Ω—ã–π –±–ª–æ–∫

    return logical_blocks


# –§—É–Ω–∫—Ü–∏—è —Å–æ—Ö—Ä–∞–Ω—è–µ—Ç —Ç–∞–±–ª–∏—Ü—É –≤ MiniO –≤ —Ñ–æ—Ä–º–∞—Ç–µ XLSX
def save_table_to_minio(bucket_name, table_name, table_data):
    """–°–æ—Ö—Ä–∞–Ω—è–µ—Ç —Ç–∞–±–ª–∏—Ü—É –≤ MinIO –≤ —Ñ–æ—Ä–º–∞—Ç–µ XLSX —Å –ø—Ä–æ–≤–µ—Ä–∫–æ–π 'Cir73SPb+' –∏ –∏–∑–º–µ–Ω–µ–Ω–∏–µ–º –≥—Ä–∞–Ω–∏—Ü/—Ü–≤–µ—Ç–∞ —Ç–µ–∫—Å—Ç–∞."""

    workbook = Workbook()
    sheet = workbook.active

    # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Å—Ç–∏–ª–∏ –¥–ª—è —è—á–µ–µ–∫
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    center_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    found_target = False  # –§–ª–∞–≥, —É–∫–∞–∑—ã–≤–∞—é—â–∏–π, –≤—Å—Ç—Ä–µ—Ç–∏–ª–æ—Å—å –ª–∏ 'Cir73SPb+'
    # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫ –∏ —Å—Ç–æ–ª–±—Ü–æ–≤
    max_rows = len(table_data)
    max_cols = max(len(row) for row in table_data)

    # –î–æ–±–∞–≤–ª—è–µ–º —Å—Ç—Ä–æ–∫–∏ —Ç–∞–±–ª–∏—Ü—ã –≤ Excel —Å –ø—Ä–∏–º–µ–Ω–µ–Ω–∏–µ–º —Å—Ç–∏–ª–µ–π
    for row_idx in range(1, max_rows + 1):
        for col_idx in range(1, max_cols + 1):
            # –ü–æ–ª—É—á–∞–µ–º –∑–Ω–∞—á–µ–Ω–∏–µ —è—á–µ–π–∫–∏ –∏–ª–∏ –æ—Å—Ç–∞–≤–ª—è–µ–º –ø—É—Å—Ç—É—é —Å—Ç—Ä–æ–∫—É
            cell_value = (
                table_data[row_idx - 1][col_idx - 1]
                if row_idx - 1 < len(table_data)
                and col_idx - 1 < len(table_data[row_idx - 1])
                else ""
            )
            cell = sheet.cell(row=row_idx, column=col_idx, value=cell_value)

            # –ï—Å–ª–∏ –≤—Å—Ç—Ä–µ—Ç–∏–ª–∏ "Cir73SPb+", –∞–∫—Ç–∏–≤–∏—Ä—É–µ–º —Ñ–ª–∞–≥
            if cell_value == "–ö–æ–Ω–µ—Ü —Ç–∞–±–ª–∏—Ü—ã, –Ω–∞—á–∞–ª–æ –ø–æ—è—Å–Ω–µ–Ω–∏–π":
                found_target = True
                cell.value = ""

            # –ü—Ä–∏–º–µ–Ω—è–µ–º —Å—Ç–∏–ª–∏
            if not found_target:
                cell.border = thin_border  # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≥—Ä–∞–Ω–∏—Ü—ã —Ç–æ–ª—å–∫–æ –¥–æ 'Cir73SPb+'
            else:
                cell.border = None  # –£–±–∏—Ä–∞–µ–º –≥—Ä–∞–Ω–∏—Ü—ã –ø–æ—Å–ª–µ 'Cir73SPb+'
                cell.font = Font(color="FFFFFF")  # –î–µ–ª–∞–µ–º —Ç–µ–∫—Å—Ç –±–µ–ª—ã–º

            # –í—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ –ø–æ —Ü–µ–Ω—Ç—Ä—É –∏ –ø–µ—Ä–µ–Ω–æ—Å —Ç–µ–∫—Å—Ç–∞
            cell.alignment = center_alignment
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≤—ã—Å–æ—Ç—É —Å—Ç—Ä–æ–∫–∏: 15 –ø–æ—Å–ª–µ –ø–æ—è–≤–ª–µ–Ω–∏—è '–ö–æ–Ω–µ—Ü —Ç–∞–±–ª–∏—Ü—ã, –Ω–∞—á–∞–ª–æ –ø–æ—è—Å–Ω–µ–Ω–∏–π'
        if found_target:
            sheet.row_dimensions[row_idx].height = 15

    # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —à–∏—Ä–∏–Ω—É –≤—Å–µ—Ö —Å—Ç–æ–ª–±—Ü–æ–≤ (50 —É—Å–ª–æ–≤–Ω—ã—Ö –µ–¥–∏–Ω–∏—Ü)
    for col in range(1, max_cols + 1):
        column_letter = sheet.cell(row=1, column=col).column_letter
        sheet.column_dimensions[column_letter].width = 50

    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –¥–∞–Ω–Ω—ã–µ –≤ –±—É—Ñ–µ—Ä –¥–ª—è XLSX
    buffer_xlsx = BytesIO()
    workbook.save(buffer_xlsx)
    buffer_xlsx.seek(0)

    # –°–æ—Ö—Ä–∞–Ω—è–µ–º XLSX –≤ MinIO
    xlsx_key = f"{minio_folder_docs_name}/{table_name}.xlsx"
    s3_client.put_object(
        Bucket=bucket_name,
        Key=xlsx_key,
        Body=buffer_xlsx,
        ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ContentDisposition="inline",  # –£–∫–∞–∑—ã–≤–∞–µ—Ç –±—Ä–∞—É–∑–µ—Ä—É –æ—Ç–∫—Ä—ã–≤–∞—Ç—å —Ñ–∞–π–ª, –∞ –Ω–µ —Å–∫–∞—á–∏–≤–∞—Ç—å
    )
    global count_table_to_save
    print(f"–¢–∞–±–ª–∏—Ü–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∞ ({count_table_to_save}) –≤ MinIO –∫–∞–∫ {table_name}.xlsx")

    count_table_to_save += 1


# –§—É–Ω–∫—Ü–∏—è —Å–æ—Ö—Ä–∞–Ω—è–µ—Ç –≤—Å–µ —Ä–∏—Å—É–Ω–∫–∏ –∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –∫–∞–∫ JPEG
def save_image_to_minio(bucket_name, image_name, image_data):
    """–°–æ—Ö—Ä–∞–Ω—è–µ—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –≤ MinIO –∫–∞–∫ JPEG —Ñ–∞–π–ª –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç –∏–º—è —Ñ–∞–π–ª–∞"""
    buffer = BytesIO()
    image_data = image_data.convert(
        "RGB"
    )  # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –≤ RGB –ø–µ—Ä–µ–¥ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ–º
    image_data.save(buffer, format="JPEG")
    s3_client.put_object(
        Bucket=bucket_name,
        Key=f"{minio_folder_docs_name}/{image_name}",
        Body=buffer.getvalue(),
        ContentType="image/jpeg",  # MIME-—Ç–∏–ø –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
        ContentDisposition="inline",  # –£–∫–∞–∑—ã–≤–∞–µ—Ç –±—Ä–∞—É–∑–µ—Ä—É –æ—Ç–∫—Ä—ã–≤–∞—Ç—å —Ñ–∞–π–ª, –∞ –Ω–µ —Å–∫–∞—á–∏–≤–∞—Ç—å
    )
    global count_image_to_save
    print(f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –∑–∞–≥—Ä—É–∂–µ–Ω–æ ({count_image_to_save}) –≤ MinIO –∫–∞–∫ {image_name}")
    count_image_to_save += 1
    return image_name  # –í–æ–∑–≤—Ä–∞—â–∞–µ–º –∏–º—è —Ñ–∞–π–ª–∞ –≤–º–µ—Å—Ç–æ —Å—Å—ã–ª–∫–∏


# –§—É–Ω–∫—Ü–∏—è –¥–ª—è –∏–∑–≤–ª–µ—á–µ–Ω–∏—è "–†–∏—Å—É–Ω–æ–∫ –•" –∏–∑ —Ç–µ–∫—Å—Ç–∞
def extract_figure_id(text):
    match = re.search(r"(–†–∏—Å—É–Ω–æ–∫ \d+)", text)
    if match:
        return match.group(1)
    return ""


# –§—É–Ω–∫—Ü–∏—è –¥–ª—è –∏–∑–≤–ª–µ—á–µ–Ω–∏—è "–¢–∞–±–ª–∏—Ü–∞ –•" –∏–∑ –Ω–∞–∑–≤–∞–Ω–∏—è —Ç–∞–±–ª–∏—Ü—ã
def extract_table_id(text):
    """
    –ò–∑–≤–ª–µ–∫–∞–µ—Ç –∏–¥–µ–Ω—Ç–∏—Ñ–∏–∫–∞—Ç–æ—Ä —Ç–∞–±–ª–∏—Ü—ã –≤ —Ñ–æ—Ä–º–∞—Ç–µ '–¢–∞–±–ª–∏—Ü–∞ X' –∏–∑ —Ç–µ–∫—Å—Ç–∞.
    –ï—Å–ª–∏ –∏–¥–µ–Ω—Ç–∏—Ñ–∏–∫–∞—Ç–æ—Ä –Ω–µ –Ω–∞–π–¥–µ–Ω, –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç –ø—É—Å—Ç—É—é —Å—Ç—Ä–æ–∫—É.
    """
    match = re.search(r"(–¢–∞–±–ª–∏—Ü–∞ \d+)", text, re.IGNORECASE)
    return match.group(1) if match else ""


# –§—É–Ω–∫—Ü–∏—è –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç Word –¥–æ–∫—É–º–µ–Ω—Ç, –∏–∑–≤–ª–µ–∫–∞—è —Ç–∞–±–ª–∏—Ü—ã, —Ç–µ–∫—Å—Ç, –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –∏ —Å–æ—Ö—Ä–∞–Ω—è—è –≤ MiniO
def extract_content_from_word(word_path, bucket_name):
    """–ò–∑–≤–ª–µ–∫–∞–µ—Ç —Ç–µ–∫—Å—Ç, —Ç–∞–±–ª–∏—Ü—ã –∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏–∑ Word —Ñ–∞–π–ª–∞, –∏–∑–±–µ–≥–∞—è –¥—É–±–ª–∏–∫–∞—Ç–æ–≤."""
    doc = Document(word_path)
    text_blocks_with_refs = []
    current_text_block = []
    current_table_data = []
    table_counter, image_counter = 1, 1
    last_was_table = False
    saved_images = set()  # –ù–∞–±–æ—Ä –¥–ª—è –æ—Ç—Å–ª–µ–∂–∏–≤–∞–Ω–∏—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–Ω—ã—Ö –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
    num_table_ff = ""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ —Ç–µ–∫—Å—Ç–∞ –∏ —Ç–∞–±–ª–∏—Ü
    for idx, block in enumerate(doc.element.body):

        if block.tag.endswith("p"):  # –û–±—Ä–∞–±–æ—Ç–∫–∞ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤
            paragraph = block.text.strip()
            if paragraph:
                if last_was_table and current_table_data:
                    # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–µ–∫—É—â—É—é —Å–æ–±—Ä–∞–Ω–Ω—É—é —Ç–∞–±–ª–∏—Ü—É –≤ MinIO –∫–∞–∫ –æ–¥–Ω—É —Ç–∞–±–ª–∏—Ü—É
                    num_table_ff = uuid.uuid4().hex[:20]
                    table_name = f"table_{num_table_ff}"
                    table_name_xlsx = f"{table_name}.xlsx"
                    save_table_to_minio(bucket_name, table_name, current_table_data)
                    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –æ–ø–∏—Å–∞–Ω–∏–µ —Ç–∞–±–ª–∏—Ü—ã
                    explanation = current_text_block[-1] if current_text_block else ""

                    table_id = extract_table_id(explanation)

                    text_blocks_with_refs.append(
                        {
                            "text": explanation,
                            "reference": table_name_xlsx,
                            "figure_id": table_id,
                            "related_table": "",
                            "origin_name_docs": name_of_origin_doc,
                        }
                    )
                    # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º —Ç–µ–∫—Å—Ç –∏–∑ —Ç–∞–±–ª–∏—Ü—ã –∏ —Å–æ—Ö—Ä–∞–Ω—è–µ–º –≤ Milvus
                    table_text_blocks = split_table_text_logically(current_table_data)
                    for block in table_text_blocks:
                        text_blocks_with_refs.append(
                            {
                                "text": block,
                                "reference": "",
                                "figure_id": "",
                                "related_table": table_name_xlsx,
                                "origin_name_docs": name_of_origin_doc,
                            }
                        )
                    current_table_data = []  # –°–±—Ä–æ—Å —Ç–µ–∫—É—â–∏—Ö –¥–∞–Ω–Ω—ã—Ö —Ç–∞–±–ª–∏—Ü—ã
                    table_counter += 1
                current_text_block.append(paragraph)
                last_was_table = False

        elif block.tag.endswith("tbl"):  # –û–±—Ä–∞–±–æ—Ç–∫–∞ —Ç–∞–±–ª–∏—Ü
            table = next(t for t in doc.tables if t._tbl == block)
            table_data = [
                [cell.text.strip() for cell in row.cells] for row in table.rows
            ]

            # –û–±—ä–µ–¥–∏–Ω—è–µ–º —Ç–∞–±–ª–∏—Ü—ã, –µ—Å–ª–∏ –æ–Ω–∏ –∏–¥—É—Ç –ø–æ–¥—Ä—è–¥
            if last_was_table:
                current_table_data.extend(table_data)
            else:
                current_table_data = table_data

            # –û–±—Ä–∞–±–æ—Ç–∫–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –≤–Ω—É—Ç—Ä–∏ —Ç–∞–±–ª–∏—Ü

            for row in table.rows:
                for cell in row.cells:
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        for run in paragraph.runs:
                            if run.element.xpath(".//a:blip"):
                                blip = run.element.xpath(".//a:blip")[0]
                                image_part = doc.part.related_parts[
                                    blip.get(
                                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                    )
                                ]
                                # –ü—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞ –¥—É–±–ª–∏–∫–∞—Ç
                                if image_part in saved_images:
                                    continue
                                saved_images.add(image_part)
                                random_suffix_image = uuid.uuid4().hex[:20]
                                image_data = Image.open(BytesIO(image_part.blob))
                                image_name = f"image_{random_suffix_image}.jpeg"
                                save_image_to_minio(bucket_name, image_name, image_data)

                                # –ü–æ–ª—É—á–∞–µ–º —Å–ª–µ–¥—É—é—â–∏–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ –¥–ª—è –æ–ø–∏—Å–∞–Ω–∏—è, –µ—Å–ª–∏ –æ–Ω —Å—É—â–µ—Å—Ç–≤—É–µ—Ç
                                if paragraph_index + 1 < len(cell.paragraphs):
                                    text_after_image = cell.paragraphs[
                                        paragraph_index + 1
                                    ].text.strip()
                                else:
                                    text_after_image = "–û–ø–∏—Å–∞–Ω–∏–µ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç"

                                # –ò–∑–≤–ª–µ–∫–∞–µ–º "–†–∏—Å—É–Ω–æ–∫ –•" –∏–∑ –æ–ø–∏—Å–∞–Ω–∏—è, –µ—Å–ª–∏ –æ–Ω–æ –µ—Å—Ç—å
                                figure_id = extract_figure_id(text_after_image)

                                # –î–æ–±–∞–≤–ª—è–µ–º –∑–∞–ø–∏—Å—å —Å `text`, `reference` –∏ `figure_id`
                                text_blocks_with_refs.append(
                                    {
                                        "text": text_after_image,
                                        "reference": image_name,
                                        "figure_id": figure_id,
                                        "related_table": f"table_{num_table_ff}.xlsx",  # –°–º–µ—â–µ–Ω–∏–µ –Ω–∞ 1 –¥–ª—è —Ç–∞–±–ª–∏—Ü—ã
                                        "origin_name_docs": name_of_origin_doc,
                                    }
                                )
                                print(
                                    f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {image_name} –∑–∞–≥—Ä—É–∂–µ–Ω–æ —Å –æ–ø–∏—Å–∞–Ω–∏–µ–º: {text_after_image}, figure_id: {figure_id}, related_table: {table_name_xlsx}"
                                )
                                image_counter += 1
            last_was_table = True

    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –≤–Ω–µ —Ç–∞–±–ª–∏—Ü
    paragraphs = iter(doc.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                # –ù–∞–π–¥–µ–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
                blip = run.element.xpath(".//a:blip")[0]
                image_part = doc.part.related_parts[
                    blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                ]
                # –ü—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞ –¥—É–±–ª–∏–∫–∞—Ç
                if image_part in saved_images:
                    continue
                saved_images.add(image_part)

                image_data = Image.open(BytesIO(image_part.blob))
                random_suffix_image = uuid.uuid4().hex[:20]
                image_name = f"image_{random_suffix_image}.jpeg"
                save_image_to_minio(bucket_name, image_name, image_data)

                # –ò—â–µ–º —Ç–µ–∫—Å—Ç –Ω–µ–ø–æ—Å—Ä–µ–¥—Å—Ç–≤–µ–Ω–Ω–æ –ø–æ—Å–ª–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                try:
                    next_paragraph = next(paragraphs)
                    text_after_image = (
                        next_paragraph.text.strip()
                        if next_paragraph.text.strip()
                        else "–û–ø–∏—Å–∞–Ω–∏–µ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç"
                    )
                except StopIteration:
                    text_after_image = "–û–ø–∏—Å–∞–Ω–∏–µ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç"

                # –ò–∑–≤–ª–µ–∫–∞–µ–º "–†–∏—Å—É–Ω–æ–∫ –•" –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –≤ –ø–æ–ª–µ figure_id
                figure_id = extract_figure_id(text_after_image)

                text_blocks_with_refs.append(
                    {
                        "text": text_after_image,
                        "reference": image_name,
                        "figure_id": figure_id,
                        "related_table": "",  # –ü–æ–ª–µ –æ—Å—Ç–∞—ë—Ç—Å—è –ø—É—Å—Ç—ã–º
                        "origin_name_docs": name_of_origin_doc,
                    }
                )
                print(
                    f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {image_name} –∑–∞–≥—Ä—É–∂–µ–Ω–æ —Å –æ–ø–∏—Å–∞–Ω–∏–µ–º: {text_after_image} –∏ 'figure_id': {figure_id}"
                )
                image_counter += 1

    return text_blocks_with_refs, " ".join(current_text_block)


# –§—É–Ω–∫—Ü–∏—è –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç –¥–∞–Ω–Ω—ã–µ –∏–∑ Word, —Å–æ–∑–¥–∞–µ—Ç —ç–º–±–µ–¥–¥–∏–Ω–≥–∏ –∏ —Å–æ—Ö—Ä–∞–Ω—è–µ—Ç –≤—Å–µ –≤ Milvus
def process_content_from_word(word_path, bucket_name):
    """–û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç —Ç–µ–∫—Å—Ç, —Ç–∞–±–ª–∏—Ü—ã –∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏–∑ Word —Ñ–∞–π–ª–∞ –∏ —Å–æ—Ö—Ä–∞–Ω—è–µ—Ç –≤ Milvus."""
    successful_embeddings_count = 0
    text_blocks_with_refs, full_text = extract_content_from_word(word_path, bucket_name)
    text_blocks = split_text_logically(full_text)

    for block in text_blocks:
        if block and block.strip():  # –ü—Ä–æ–≤–µ—Ä–∫–∞, —á—Ç–æ–±—ã –±–ª–æ–∫ —Ç–µ–∫—Å—Ç–∞ –Ω–µ –±—ã–ª –ø—É—Å—Ç—ã–º
            embedding = create_embeddings(block)
            if embedding is None:
                continue
            embedding_np = np.array(embedding, dtype=np.float32).tolist()
            data = [[embedding_np], [block], [""], [""], [""], [name_of_origin_doc]]
            collection.insert(data)
            successful_embeddings_count += 1
            global count_embedding_save
            print(
                f"–≠–º–±–µ–¥–¥–∏–Ω–≥ –∏ —Ç–µ–∫—Å—Ç —É—Å–ø–µ—à–Ω–æ ({count_embedding_save})–¥–æ–±–∞–≤–ª–µ–Ω—ã –¥–ª—è –±–ª–æ–∫–∞ {successful_embeddings_count}."
            )
            count_embedding_save += 1
        else:
            print("–ü—É—Å—Ç–æ–π —Ç–µ–∫—Å—Ç, –ø—Ä–æ–ø—É—Å–∫ —ç–º–±–µ–¥–¥–∏–Ω–≥–∞")

    for ref_info in text_blocks_with_refs:
        text = ref_info["text"]
        reference = ref_info["reference"]
        figure_id = ref_info["figure_id"]
        related_table = ref_info["related_table"]
        origin_name_docs = ref_info["origin_name_docs"]
        if text and text.strip():  # –ü—Ä–æ–≤–µ—Ä–∫–∞, —á—Ç–æ–±—ã —Ç–µ–∫—Å—Ç –æ–ø–∏—Å–∞–Ω–∏—è –Ω–µ –±—ã–ª –ø—É—Å—Ç—ã–º
            embedding = create_embeddings(text)
            if embedding is None:
                continue
            embedding_np = np.array(embedding, dtype=np.float32).tolist()
            data = [
                [embedding_np],
                [text],
                [reference],
                [figure_id],
                [related_table],
                [origin_name_docs],
            ]
            collection.insert(data)
            successful_embeddings_count += 1
            print(
                f"–≠–º–±–µ–¥–¥–∏–Ω–≥ –∏ –ø–æ—è—Å–Ω–µ–Ω–∏–µ —É—Å–ø–µ—à–Ω–æ –¥–æ–±–∞–≤–ª–µ–Ω—ã –¥–ª—è –æ–±—ä–µ–∫—Ç–∞: –†–µ—Ñ–µ—Ä–µ–Ω—Å - '{reference}', –†–æ–¥–∏—Ç–µ–ª—å—Å–∫–∏–π —Ñ–∞–π–ª - '{related_table}'"
            )
            count_embedding_save += 1
        else:
            print("–ü—É—Å—Ç–æ–µ –æ–ø–∏—Å–∞–Ω–∏–µ, –ø—Ä–æ–ø—É—Å–∫ —ç–º–±–µ–¥–¥–∏–Ω–≥–∞ –¥–ª—è –æ–±—ä–µ–∫—Ç–∞:", reference)

    collection.flush()
    print("–í—Å–µ —ç–º–±–µ–¥–¥–∏–Ω–≥–∏ –∏ –¥–∞–Ω–Ω—ã–µ —É—Å–ø–µ—à–Ω–æ –¥–æ–±–∞–≤–ª–µ–Ω—ã –≤ Milvus.")
    print(f"–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —É—Å–ø–µ—à–Ω–æ —Å–æ–∑–¥–∞–Ω–Ω—ã—Ö —ç–º–±–µ–¥–¥–∏–Ω–≥–æ–≤: {successful_embeddings_count}")


# –ü—Ä–∏–º–µ—Ä –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è
word_path = path_of_doc_for_convert
process_content_from_word(word_path, name_of_bucket_minio)

# –°–æ–∑–¥–∞–Ω–∏–µ –∏ –∑–∞–≥—Ä—É–∑–∫–∞ –∏–Ω–¥–µ–∫—Å–∞ –≤ Milvus
index_params = {"index_type": "IVF_FLAT", "metric_type": "L2", "params": {"nlist": 128}}
collection.create_index(field_name="embedding", index_params=index_params)
collection.load()

print(
    f"–ò–Ω–¥–µ–∫—Å —É—Å–ø–µ—à–Ω–æ —Å–æ–∑–¥–∞–Ω –∏ –∫–æ–ª–ª–µ–∫—Ü–∏—è '{collection_name}' –∑–∞–≥—Ä—É–∂–µ–Ω–∞ –≤ –ë–î '{change_db_of_milvus}'."
)
